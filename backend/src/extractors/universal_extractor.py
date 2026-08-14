import pdfplumber
import os
import re
from openai import OpenAI
from src.schemas.master_schema import MasterSBCSchema
from src.extractors.text_cleaner import TextCleaner
from src.extractors.ocr_helper import ocr_pdf_header, ocr_image, ocr_pdf_full
try:
    import docx
except ImportError:
    docx = None

class UniversalExtractor:
    def __init__(self):
        # Initialize OpenAI client with extended timeouts to handle network issues
        self.client = OpenAI(
            timeout=120.0,  # 120 second timeout for API calls
            max_retries=3   # Retry up to 3 times on failure
        )
        self.cleaner = TextCleaner()

    def _extract_in_network_value(self, raw_value: str) -> str:
        if raw_value is None:
            return None
        raw_value = raw_value.strip()
        if raw_value == "":
            return None

        # Prefer explicit In-Network values if both In-Network and Out-of-Network values are present.
        in_network_match = re.search(
            r'(?:In[- ]Network|In Network|In-Network)\s*[:\-]?\s*([^;|\n]+)',
            raw_value,
            re.IGNORECASE,
        )
        if in_network_match:
            return in_network_match.group(1).strip()

        if re.search(r'Out[- ]?of[- ]?Network', raw_value, re.IGNORECASE):
            # Prefer text before the out-of-network section if it appears first.
            parts = re.split(r'Out[- ]?of[- ]?Network', raw_value, flags=re.IGNORECASE)
            if parts and parts[0].strip():
                return parts[0].strip()

        return raw_value

    def _extract_non_preferred_cost(self, text: str) -> str:
        """Return the cost tied to an explicit Non-preferred label, if present."""
        if not text:
            return None
        patterns = [
            r'([\$][\d,]+(?:\.\d+)?|\d+%)\s*\(\s*non[- ]?preferred\s*\)',
            r'non[- ]?preferred\s*[:\-]?\s*([\$][\d,]+(?:\.\d+)?|\d+%)',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1)
        return None

    def _normalize_pharmacy_tier(self, raw_value: str) -> str:
        if raw_value is None:
            return None
        raw_value = raw_value.strip()
        if raw_value == "":
            return None

        raw_value_lower = raw_value.lower()
        if 'no charge' in raw_value_lower or 'not applicable' in raw_value_lower:
            return '$0'

        raw_value = self._extract_in_network_value(raw_value)

        raw_value_lower_full = raw_value.lower()
        has_preferred_non_preferred_pair = (
            re.search(r'non[- ]?preferred', raw_value_lower_full)
            and re.search(r'preferred', raw_value_lower_full)
        )
        if has_preferred_non_preferred_pair:
            non_preferred_cost = self._extract_non_preferred_cost(raw_value)
            if non_preferred_cost:
                return non_preferred_cost

        split_pattern = r';|\band\b|\bor\b|\||/' if has_preferred_non_preferred_pair else r';|\band\b|\bor\b|\|'

        # Prefer the retail value if both retail and mail-order values are present.
        parts = re.split(split_pattern, raw_value, flags=re.IGNORECASE)

        def _is_non_preferred_part(part_lower: str) -> bool:
            return bool(re.search(r'non[- ]?preferred', part_lower))

        def _is_preferred_only_part(part_lower: str) -> bool:
            return 'preferred' in part_lower and not _is_non_preferred_part(part_lower)

        # FIRST PASS: Explicit Preferred vs Non-preferred retail — take Non-preferred only
        # when BOTH labels appear (does not affect specialty_rx_description or other fields).
        has_non_preferred = any(_is_non_preferred_part(p.strip().lower()) for p in parts)
        has_preferred_only = any(_is_preferred_only_part(p.strip().lower()) for p in parts)
        non_preferred_part = None
        if has_non_preferred and has_preferred_only:
            for part in parts:
                part_text = part.strip()
                if _is_non_preferred_part(part_text.lower()):
                    non_preferred_part = part_text
                    break

        if non_preferred_part:
            raw_value = non_preferred_part
        else:
            # SECOND PASS: Prefer "Participating" (excluding "Preferred Participating")
            # This handles cases where both "Preferred Participating" and "Participating" values exist
            # We want pure "Participating" over "Preferred Participating" variant
            participating_part = None
            for part in parts:
                part_text = part.strip()
                part_lower = part_text.lower()
                if 'participating' in part_lower and 'preferred' not in part_lower:
                    participating_part = part_text
                    break

            if participating_part:
                raw_value = participating_part
            else:
                retail_part = None
                fallback_part = None
                for part in parts:
                    part_text = part.strip()
                    part_lower = part_text.lower()
                    if 'retail' in part_lower and 'mail' not in part_lower:
                        retail_part = part_text
                        break
                    if fallback_part is None and 'mail' not in part_lower:
                        fallback_part = part_text
                if retail_part:
                    raw_value = retail_part
                elif fallback_part:
                    raw_value = fallback_part

        # Remove descriptive parentheticals like (retail) or (mail order)
        raw_value = re.sub(r'\s*\([^)]*\)', '', raw_value).strip()

        # If there is a direct dollar amount, return only that
        dollar_match = re.search(r'\$[\d,]+(?:\.\d+)?', raw_value)
        if dollar_match:
            return dollar_match.group(0)

        # Preserve simple percentage expressions when there is no dollar amount
        percent_match = re.search(r'\d+%', raw_value)
        if percent_match:
            return percent_match.group(0)

        return raw_value

    def _percent_near_keyword(self, text: str, keyword_pattern: str, window: int = 80) -> bool:
        """Return True if a percentage (e.g. '20%') or the word 'percent' appears
        within `window` characters of any occurrence of `keyword_pattern` in text.
        This helps avoid assigning a coinsurance value to a visit row when the
        percent actually refers to nearby imaging/facility rows.
        """
        if not text:
            return False
        flags = re.IGNORECASE
        for m in re.finditer(keyword_pattern, text, flags):
            start, end = m.start(), m.end()
            lo = max(0, start - window)
            hi = min(len(text), end + window)
            snippet = text[lo:hi]
            if re.search(r'\d+%|percent', snippet, flags):
                return True
        return False

    def _needs_ocr_header(self, raw_text: str) -> bool:
        # If the text contains a header line beginning with a leading colon, the carrier prefix may be missing.
        if re.search(r'^\s*:\s*[^\n]+plan\s*type\s*:', raw_text, re.IGNORECASE | re.MULTILINE):
            return True

        lines = [line for line in raw_text.splitlines() if line.strip()]
        if len(lines) >= 2 and re.match(r'^\s*:\s*', lines[1]):
            return True

        return False

    def _has_cid_encoding(self, text: str) -> bool:
        """Detect PDFs where embedded fonts lack Unicode mapping (pdfplumber outputs (cid:NNN))."""
        cid_tokens = re.findall(r'\(cid:\d+\)', text)
        if len(cid_tokens) < 10:
            return False
        stripped = re.sub(r'\(cid:\d+\)', '', text)
        readable_words = len(re.findall(r'[a-zA-Z]{3,}', stripped))
        return len(cid_tokens) >= max(20, readable_words * 3)

    def _format_ocr_snippet(self, ocr_text: str, max_length: int = 180) -> str:
        snippet = ' '.join(ocr_text.split())
        if len(snippet) > max_length:
            snippet = snippet[:max_length].rstrip() + '...'
        return snippet

    def _extract_plan_name_from_text(self, raw_text: str, extracted_name: str, extracted_carrier: str) -> str:
        candidate = extracted_name
        COV_FOR = r'coverage\s*for\s*:'
        PLAN_TYPE = r'plan\s*type\s*:'

        def is_date_range(text: str) -> bool:
            return bool(re.search(r'\d{1,2}/\d{1,2}/\d{2,4}\s*-\s*\d{1,2}/\d{1,2}/\d{2,4}', text))

        def clean_metadata(text: str) -> str:
            """Clean common metadata labels out of the plan name candidate."""
            t = re.sub(rf'{COV_FOR}.*', '', text, flags=re.IGNORECASE)
            # Handles PDF line-breaks where "Coverage" ends one line and "for: ..." starts the next
            t = re.sub(r'\s+for\s*:.*', '', t, flags=re.IGNORECASE)
            t = re.sub(rf'{PLAN_TYPE}.*', '', t, flags=re.IGNORECASE)
            t = re.sub(r'Summary of Benefits.*', '', t, flags=re.IGNORECASE)
            return t.strip(' |,-/\t:')

        def strip_page_marker(text: str) -> str:
            return re.sub(r'^---\s*PAGE\s*\d+\s*(?:TEXT|STRUCTURED).*?--+\s*', '', text, flags=re.IGNORECASE).strip()

        plan_name = None
        plan_name_from_header = False  # True when plan_name came from the authoritative SBC header line
        lines = raw_text.splitlines()

        # ── Priority 0: UnitedHealthcare Specific Pattern ─────────────────────
        # Look for "SBC_filename" pattern or title line with plan details
        # Examples: "SBC_120-25_UnitedHealthcare_Choice_Plus_HDHP_3300"
        #           "UnitedHealthcare Choice Plus HDHP 3300"
        # ────────────────────────────────────────────────────────────────────────
        
        # First try: filename pattern in header
        uhc_filename_pattern = re.search(
            r'SBC[_\-]\d+[_\-]\d+[_\-]UnitedHealthcare[_\s]*([A-Za-z0-9\s_]+?)(?:\s+\d+)?(?:\.pdf)?',
            raw_text, 
            re.IGNORECASE
        )
        if uhc_filename_pattern:
            uhc_candidate = uhc_filename_pattern.group(1).strip()
            # Clean up underscores and normalize spaces
            uhc_candidate = re.sub(r'[_]+', ' ', uhc_candidate).strip()
            uhc_candidate = re.sub(r'\s+', ' ', uhc_candidate)
            # Remove trailing page numbers
            uhc_candidate = re.sub(r'\s+\d+\s*$', '', uhc_candidate).strip()
            if uhc_candidate and len(uhc_candidate) > 3:
                plan_name = f"UnitedHealthcare {uhc_candidate}"
                print(f"    [PLAN_NAME] UHC filename pattern found: {plan_name}")

        # Second try: direct text line pattern  
        if not plan_name:
            uhc_direct_pattern = re.search(
                r'UnitedHealthcare\s+([A-Za-z0-9\s]+(?:HDHP|PPO|HMO|Choice|Options)[A-Za-z0-9\s]*)',
                raw_text, 
                re.IGNORECASE
            )
            if uhc_direct_pattern:
                uhc_candidate = uhc_direct_pattern.group(1).strip()
                uhc_candidate = re.sub(r'\s+', ' ', uhc_candidate)
                # Remove trailing page numbers or artifacts
                uhc_candidate = re.sub(r'\s+\d+\s*$', '', uhc_candidate).strip()
                if uhc_candidate and len(uhc_candidate) > 3:
                    plan_name = f"UnitedHealthcare {uhc_candidate}"
                    print(f"    [PLAN_NAME] UHC direct pattern found: {plan_name}")

        # Third try: UHC header with collapsed spacing (e.g. ":ChoicePlusDY1P/L27S Coveragefor:...")
        if not plan_name:
            uhc_choice_line = re.search(
                r':?\s*(Choice\s*Plus\s*[A-Za-z0-9/\-]+)\s*Coverage\s*for\s*:',
                raw_text,
                re.IGNORECASE,
            )
            if uhc_choice_line:
                uhc_candidate = re.sub(r'\s+', ' ', uhc_choice_line.group(1).strip())
                plan_name = f"UnitedHealthcare {uhc_candidate}"
                print(f"    [PLAN_NAME] UHC Choice Plus header found: {plan_name}")

        # ── Priority 1 ─────────────────────────────────────────────────────────
        # Find the SBC header line containing BOTH "Coverage for:" and "Plan Type:".
        #
        #   Pattern A (Cigna):  GALA CORPORATION: $3,000 COPAY Coverage for: Ind/Family | Plan Type: PPO
        #                       → extract text between employer colon and "Coverage for:"
        #
        #   Pattern B (WellMark): Coverage for: Single & Family | Plan Type: PPO
        #                         EnhancedBlueSM 1500 PPO   ← plan name on NEXT line
        # ────────────────────────────────────────────────────────────────────────
        if not plan_name:
            for idx, line in enumerate(lines):
                if re.search(COV_FOR, line, re.IGNORECASE) and re.search(PLAN_TYPE, line, re.IGNORECASE):

                    # --- Pattern A: text BEFORE "Coverage for:" on this same line ---
                    before_cov = re.split(COV_FOR, line, maxsplit=1, flags=re.IGNORECASE)[0].strip()
                    before_cov = strip_page_marker(before_cov)
                    before_cov = re.sub(r'^Summary of Benefits[^:]*:?', '', before_cov, flags=re.IGNORECASE).strip()
                    before_cov = re.sub(r'Coverage Period\s*:.*$', '', before_cov, flags=re.IGNORECASE).strip(' |,-/\t')

                    if ':' in before_cov:
                        # "EMPLOYER_NAME: PLAN_NAME" → take part after the LAST colon
                        after_colon = before_cov.rsplit(':', 1)[1].strip()
                        cleaned = clean_metadata(after_colon)
                        if cleaned and len(cleaned) > 2 and not is_date_range(cleaned):
                            plan_name = cleaned
                            plan_name_from_header = True

                    # --- Pattern B: plan name on the NEXT non-metadata line ---
                    if not plan_name:
                        skip_phrases = [
                            'coverage for:', 'plan type:', 'summary of benefits',
                            'the summary of', 'this is only a summary', 'share the cost',
                            'note:', 'important questions', 'what is the overall',
                            '1-800', '1-888', 'all copayment', 'all coinsurance',
                            # OCR scanned PDFs: glossary disclaimer runs right after the header line
                            'allowed amount', 'balance billing', 'underlined terms',
                            'see the glossary', 'for general definitions', 'copayment, deductible',
                        ]
                        for next_line in lines[idx + 1:idx + 6]:
                            nl_clean = next_line.strip()
                            if not nl_clean:
                                continue
                            if any(kw in nl_clean.lower() for kw in skip_phrases):
                                continue
                            if is_date_range(nl_clean):
                                continue
                            nl_clean = strip_page_marker(nl_clean)
                            cleaned = clean_metadata(nl_clean)
                            if cleaned and len(cleaned) > 3:
                                plan_name = cleaned
                                plan_name_from_header = True
                                break
                    break  # Only check the first occurrence of this header line

        # ── Priority 2 ─────────────────────────────────────────────────────────
        # Text BEFORE "Coverage Period:" on the same line (e.g., "BlueOptions 05906 Coverage Period:...")
        # ────────────────────────────────────────────────────────────────────────
        if not plan_name:
            pre_coverage_match = re.search(r'^\s*(.*?)\s*Coverage Period:', raw_text, re.IGNORECASE | re.MULTILINE)
            if pre_coverage_match:
                pc_candidate = strip_page_marker(pre_coverage_match.group(1).strip())
                cleaned = clean_metadata(pc_candidate)
                if cleaned and len(cleaned) > 3 and not is_date_range(cleaned):
                    plan_name = cleaned

        # ── Priority 3 ─────────────────────────────────────────────────────────
        # Lines AFTER "Coverage Period:" — 8-line lookahead to reach past metadata lines.
        # ────────────────────────────────────────────────────────────────────────
        if not plan_name:
            skip_phrases_p3 = [
                'summary of benefits', 'share the cost', 'this is only a summary',
                'note:', 'for general definitions', 'the summary of',
            ]
            for idx, line in enumerate(lines):
                if 'coverage period:' in line.lower():
                    for next_line in lines[idx + 1:idx + 8]:
                        nl_clean = next_line.strip()
                        if not nl_clean:
                            continue
                        nl_lower = nl_clean.lower()
                        if any(kw in nl_lower for kw in skip_phrases_p3):
                            continue
                        if is_date_range(nl_clean):
                            continue
                        nl_clean = strip_page_marker(nl_clean)
                        cleaned = clean_metadata(nl_clean)
                        # Skip lines that are purely metadata (e.g. "Coverage for: Ind/Family | Plan Type: PPO")
                        if not cleaned or len(cleaned) <= 3:
                            continue
                        if re.fullmatch(r'coverage\s*for\s*:.*', nl_clean, re.IGNORECASE):
                            continue
                        plan_name = cleaned
                        break
                    break

        # ── Shared rejection rules (apply to BOTH regex result AND LLM candidate) ─
        BAD_PHRASES = [
            'your rights', 'there are agencies', 'contact information', 'this is only a summary',
            'welcometouhc.com', 'for general definitions', 'allowed amount', 'balance billing',
            'coinsurance, copayment', 'underlined terms', 'see the glossary', 'copayment, deductible',
        ]
        # Employer/group name suffixes — a plan name never ends with these
        EMPLOYER_SUFFIXES = [
            r'\bINC\.?\b', r'\bLLC\.?\b', r'\bPEO\b', r'\bCORP\.?\b',
            r'\bGROUP\b', r'\bSERVICES\b', r'\bSOLUTIONS\b', r'\bHOLDINGS\b',
            r'\bTRUST\b', r'\bASSOCIATION\b', r'\bFOUNDATION\b', r'\bENTERPRISES\b',
        ]

        def _is_bad_plan_name(name: str) -> bool:
            if not name:
                return True
            nl = name.lower()
            if any(p in nl for p in BAD_PHRASES):
                return True
            if any(re.search(pat, name, re.IGNORECASE) for pat in EMPLOYER_SUFFIXES):
                return True
            return False

        # Apply rejection to regex-extracted plan_name
        if plan_name and (_is_bad_plan_name(plan_name) or is_date_range(plan_name)):
            print(f"    [PLAN_NAME] Rejected regex result (employer/glossary): {plan_name!r}")
            plan_name = None

        # Apply rejection to LLM candidate
        if candidate:
            candidate = clean_metadata(candidate.strip())
            if (is_date_range(candidate)
                    or candidate.lower() == extracted_carrier.strip().lower()
                    or candidate.lower() == 'plan'
                    or len(candidate.split()) > 12
                    or _is_bad_plan_name(candidate)):
                candidate = None
                print(f"    [PLAN_NAME] Rejected LLM candidate: contains bad phrases or invalid format")

        if candidate and plan_name:
            # If the regex result came from the authoritative SBC header line, always prefer it
            if plan_name_from_header:
                print(f"    [PLAN_NAME] Using header result (authoritative): {plan_name}")
                return plan_name
            # Otherwise prefer regex-detected plan_name if LLM candidate is shorter/less specific
            if len(candidate.split()) < max(4, len(plan_name.split())):
                print(f"    [PLAN_NAME] Using regex result: {plan_name}")
                return plan_name
            print(f"    [PLAN_NAME] Using LLM result: {candidate}")
            return candidate

        final_name = candidate or plan_name or extracted_carrier
        print(f"    [PLAN_NAME] Final result: {final_name}")
        return final_name

    def _extract_pdf(self, file_path: str) -> str:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages) if pdf.pages else 1
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t:
                    text += f"--- PAGE {i+1} TEXT ---\n" + t + "\n"
                
                tables = page.extract_tables()
                if tables:
                    text += f"--- PAGE {i+1} STRUCTURED TABLES ---\n"
                    for table in tables:
                        for row in table:
                            clean_row = [str(cell).replace('\n', ' ').strip() if cell else '' for cell in row]
                            text += " | ".join(clean_row) + "\n"
                    text += "\n"

        # CID placeholders mean broken font encoding — OCR reads the rendered page instead.
        if self._has_cid_encoding(text):
            cid_count = len(re.findall(r'\(cid:\d+\)', text))
            print(
                f"  [OCR] CID font encoding detected ({cid_count} tokens). "
                f"Attempting full OCR on {os.path.basename(file_path)}..."
            )
            text = ocr_pdf_full(file_path)
        else:
            # If extracted text is very short or has low text density, it's likely a scanned PDF.
            text_for_density = re.sub(r'\(cid:\d+\)', ' ', text)
            alnum_count = len(re.findall(r'[a-zA-Z0-9]', text_for_density))
            avg_alnum_per_page = alnum_count / num_pages
            if alnum_count < 150 or avg_alnum_per_page < 60:
                print(
                    f"  [OCR] Low text volume detected ({alnum_count} chars, "
                    f"avg {avg_alnum_per_page:.1f}/page). Attempting full OCR on {os.path.basename(file_path)}..."
                )
                text = ocr_pdf_full(file_path)

        if self._needs_ocr_header(text):
            ocr_text = ocr_pdf_header(file_path)
            if ocr_text:
                snippet = self._format_ocr_snippet(ocr_text)
                print(f"  [OCR] Header fallback used; OCR snippet: {snippet}")
                text = ocr_text + "\n" + text
        return text

    def _extract_docx(self, file_path: str) -> str:
        if not docx:
            print("  [ERR] python-docx not installed. Skipping Word file.")
            return ""
        print("  [Extraction] Using: python-docx (Word Parsing)")
        doc = docx.Document(file_path)
        text = ""
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text += para.text + "\n"
        
        for i, table in enumerate(doc.tables):
            text += f"--- TABLE {i+1} STRUCTURED ---\n"
            for row in table.rows:
                row_text = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                text += " | ".join(row_text) + "\n"
            text += "\n"
        return text

    def _extract_image(self, file_path: str) -> str:
        print("  [Extraction] Using: OCR (Tesseract) for Image")
        print(f"  [OCR] Processing image {os.path.basename(file_path)}...")
        return ocr_image(file_path)

    def extract_text(self, file_path: str, save_raw_path: str = None, filename: str = None) -> MasterSBCSchema:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Could not find {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        text = ""
        
        if ext == '.pdf':
            text = self._extract_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            text = self._extract_docx(file_path)
        elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            text = self._extract_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        if not text.strip():
            raise ValueError(f"No text could be extracted from {file_path}")

        # Clean the raw text to remove noise before sending to AI
        text = self.cleaner.clean(text)

        # Save raw extracted text if path is provided
        if save_raw_path:
            os.makedirs(os.path.dirname(save_raw_path), exist_ok=True)
            with open(save_raw_path, 'w', encoding='utf-8') as f:
                f.write(f"=== RAW TEXT EXTRACTED FROM: {os.path.basename(file_path)} ===\n")
                f.write(f"=== Total characters: {len(text)} ===\n\n")
                f.write(text)
            print(f"[OK] Raw text saved: {save_raw_path}")

        FORMATTING_RULES = """
STRICT VALUE FORMATTING RULES - Follow these exactly:
1. DOLLAR VALUES: Keep the $ symbol. Strip all trailing words. Example: "$250 copay per visit" -> "$250". "$10 copay" -> "$10".
2. PERCENTAGE VALUES: Keep the % symbol. Strip all trailing words. Example: "20% coinsurance" -> "20%". "40% after deductible" -> "40%".
3. NO CHARGE: Convert "No Charge", "no charge", or "Not Applicable" to "$0".
4. PHARMACY DEDUCTIBLES with Individual/Family: Return ONLY the Individual value. Example: "$150 Individual/$450 Family" -> "$150".
5. MIXED DOLLAR+PERCENT (e.g. "$200 or 20%"): Return as-is: "$200 or 20%".
6. NOT COVERED: If any COPAY field is "Not Covered", return "$0". If any COINSURANCE field is "Not Covered", return "0%". DO NOT return "Not Covered" as a value for copay or coinsurance fields.
7. DO NOT include words like "copay", "coinsurance", "per visit", "per stay", "after deductible".
8. COINSURANCE FIELDS: These expect a % value. Example: "20% coinsurance" -> "20%".
9. COPAY FIELDS: These expect a $ value. Example: "$35 copay" -> "$35".
10. DEDUCTIBLE AND OOP MAX FIELDS (individual_deductible, family_deductible, individual_oop_max, family_oop_max): Keep the $ symbol and comma. Example: "$1,500". SPECIAL RULE FOR OUT-OF-NETWORK OOP MAX ONLY: If a specific dollar value is present, extract it, BUT if the value is greater than $99,999 (e.g., $100,000), you MUST cap it and return "$99,999". If the word "Unlimited" is EXPLICITLY present, return "$99,999". If Out-of-Network is not covered/applicable/mentioned, return "$0". DO NOT return "$99,999" unless the value exceeds $99,999 or "Unlimited" is explicitly written.
11. 100% MINUS RULE (PATIENT RESPONSIBILITY): If the document lists what the PLAN pays for coinsurance (e.g., "Plan pays 80%", "Plan pays 70%"), or if any extracted coinsurance is between 60% and 100% (inclusive), you MUST subtract it from 100% to calculate the patient's responsibility. For example, if the document says "80%", extract it as "20%". If it says "70%", extract it as "30%". CRITICAL: The extracted coinsurance MUST be the PATIENT responsibility.

CRITICAL: COPAY vs COINSURANCE DETECTION RULES:
- IF the document shows ONLY percentage values (e.g., "10% coinsurance") for a service, put the percentage in COINSURANCE field and "$0" in COPAY field.
- IF the document shows ONLY dollar amounts (e.g., "$35 copay") for a service, put the dollar amount in COPAY field and "0%" in COINSURANCE field.
- IF the document shows BOTH (e.g., "$35 copay, 20% coinsurance"), extract BOTH values into their respective fields.
- IN-NETWORK FIRST PRIORITY: When a service shows two columns - "Participating Provider (In-Network)" and "Non-Participating Provider (Out-of-Network)" - ALWAYS extract from the IN-NETWORK column FIRST. The In-Network column is the correct/primary value. Example: If "Primary care" shows "No Charge/visit; deductible does not apply" in In-Network and "50% coinsurance" in Out-of-Network, extract ONLY "No Charge" from In-Network. Do NOT extract the Out-of-Network value.
- PIPE SEPARATED TABLES CAUTION: If the text uses pipe separators like "| $40 copayment | 40% coinsurance |", the pipes separate columns! The first column is In-Network and the second is Out-of-Network. You MUST COMPLETELY IGNORE the Out-of-Network value. In this example, extract "$40" for copay and "0%" for coinsurance. Do NOT extract the 40% coinsurance because it belongs to the Out-of-Network column.
- MULTI-COLUMN TABLES (Anthem/3+ columns): CRITICAL - Some insurance documents (especially Anthem) have a "Level 1 Pharmacy-Only" column BEFORE the "In-Network Provider" column. ALWAYS extract from the "In-Network Provider" column. NEVER extract from "Level 1 Pharmacy" or "Out-of-Network" columns. Example: If a row shows "| $20 | $30 | $50 |", extract ONLY $30.
- ANTHEM "NOT APPLICABLE" EXCEPTION: If you see the phrase "Not Applicable" before an In-Network percentage for a medical service (e.g., "| Not Applicable | 45% coinsurance | 50% coinsurance |"), you MUST COMPLETELY IGNORE "Not Applicable" and extract the In-Network percentage (45%). Do NOT convert "Not Applicable" to $0 for Medical Services in this scenario.
- IF YOU SEE COINSURANCE IN BOTH COLUMNS (e.g., "| 20% | 45% |" where first is Level 1 and second is In-Network), extract ONLY the In-Network value (45%). This happens frequently in Anthem plans.
- NEVER put percentage values in copay fields or dollar values in coinsurance fields.
- ACCESS+ SPECIALIST PRIORITY: If the "Specialist visit" row lists BOTH "Access+ Specialist" and "Other Specialist" costs in the In-Network column (e.g., "Access+ Specialist: $50/visit / Other Specialist: $25/visit"), you MUST ALWAYS extract the "Access+ Specialist" value ($50) into the specialist_copay field. Completely ignore the "Other Specialist" value. This rule applies exclusively to the specialist_copay and specialist_coinsurance fields. Example: "Access+ Specialist: $50/visit Other Specialist: $25/visit" → extract "$50" for specialist_copay.
- PREFERRED VS PARTICIPATING MEDICAL PRIORITY: For every medical field (Primary Care, Specialist, Lab, X-Ray, Imaging, Inpatient, Outpatient, Emergency, Urgent Care, etc.), if the Network Provider (In-Network) cell contains BOTH a "Preferred" value AND a "Participating" value, you MUST ALWAYS select and extract the "Participating" value. Completely ignore the "Preferred" value. This applies regardless of whether the value is a copay or coinsurance. Examples: "20% coinsurance for Preferred / 40% coinsurance for Participating" → extract "40% coinsurance". "$20 copay (Preferred) / $40 copay (Participating)" → extract "$40". This rule ONLY activates when both Preferred and Participating are explicitly present in the same cell. If the cell contains only one value with no Preferred/Participating split, extract that single value normally using existing logic.


MODIFIER FIELD RULES:
11. MEDICAL MODIFIER FIELDS (e.g., primary_care_copay_modifier, inpatient_coinsurance_modifier): 
    - If text says "deductible does not apply", "deductible doesn't apply", or "no charge" → return "Deductible Waived"
    - If text says "after deductible" or shows deductible applies → return "After Deductible"  
    - Otherwise, return null
12. DEDUCTIBLE STATUS FIELDS (e.g., primary_care_copay_deductible_status, primary_care_coinsurance_deductible_status):
    - If the text explicitly says "deductible does not apply", "no charge", or "after deductible", extract that EXACT phrase into the deductible_status field. Otherwise return null.
    - CRITICAL ANTI-BLEEDING RULE: Do NOT extract "deductible does not apply" for a service (like Lab Services) unless that exact phrase is physically located inside that specific service's row in the document. Do not copy it from adjacent rows like Primary Care or Preventive Care. If it is not explicitly stated in that specific row, return null.
13. PHARMACY MODIFIER FIELDS (e.g., tier_1_copay_modifier, tier_2_coinsurance_modifier): 
    - If text says "deductible does not apply" for that tier → return "Rx - Deductible Waived"
    - If text mentions "rx deductible" or "prescription deductible" → return "Rx - After Rx Deductible"
    - If text mentions "plan deductible" → return "Rx - After Plan Deductible"
    - Otherwise, return null

HOSPITAL SURGICAL EXTRACTION RULES (FACILITY-ONLY):
13. INPATIENT HOSPITAL COPAY & COINSURANCE: Extract ONLY from the "Facility" row (e.g., "Facility fee (e.g., ambulatory surgery center)" or "Facility fee (e.g., hospital room)") under the "If you have a hospital stay" section. Do NOT extract from "Physician/surgeon fees" row or any other non-facility rows. These values go into inpatient_copay and inpatient_coinsurance fields. CRITICAL: Only extract inpatient facility fees when the section is explicitly labeled for hospital/inpatient services. Do not extract facility fees from emergency care, urgent care, or other non-hospital sections. HOSPITAL PRIORITY: If the row lists multiple facility types (e.g., "Ambulatory Surgery Center" and "Hospital"), always extract the "Hospital" value.
14. OUTPATIENT HOSPITAL COPAY & COINSURANCE: Extract ONLY from the "Facility" row under the "If you have outpatient surgery" or "outpatient hospital" sections (e.g., "Facility fee (e.g., ambulatory surgery center)"). Do NOT extract from "Physician/surgeon fees" row or facility fees from emergency/urgent care sections. These values go into op_hospital_copay and op_hospital_coinsurance fields. CRITICAL: Facility fees for outpatient services must be from the outpatient surgery/hospital section ONLY, not from emergency or urgent care rows. HOSPITAL PRIORITY: If the facility row lists multiple types (e.g., "Ambulatory Surgery Center: $300" and "Hospital: $400"), always extract the "Hospital" value.
15. CRITICAL: Physician/surgeon fees are separate services and must NOT be confused with facility fees. Look for explicit "Facility" or "Facility fee" labels in the table row. If only physician/surgeon fees exist without facility fees, extract "$0" for facility-based fields. Do NOT extract facility fees from emergency medical transportation, emergency room, or urgent care sections even if they mention "facility" — these belong to different service categories and should not populate inpatient_copay, inpatient_coinsurance, op_hospital_copay, or op_hospital_coinsurance fields.

LAB AND X-RAY EXTRACTION RULES (SPLIT LAYOUT HANDLING):
15a. LAB AND X-RAY COINSURANCE: When Lab and X-ray appear on SEPARATE LINES with percentages like "Lab: 20% coinsurance" and "X-ray: 20% coinsurance", extract BOTH values from their respective lines. Do NOT treat them as missing values just because they appear on different lines from the "Diagnostic test" label. Extract from the IN-NETWORK column ONLY (ignore "Not Covered" or Out-of-Network values that appear after a pipe |).
    CRITICAL HOSPITAL/FACILITY PRIORITY FOR LAB AND X-RAY: Some documents use the word "Facility" instead of "Hospital" to indicate the same thing. These two terms are equivalent — treat "Lab Facility" the same as "Hospital Lab", and "Radiology Facility" the same as "Hospital X-ray". If the document lists multiple Lab sub-types (e.g., "Lab Office", "Lab Facility", "Free Standing Lab", "Hospital Lab"), you MUST extract ONLY the "Facility" or "Hospital" sub-type and completely ignore the "Office" and "Free Standing" sub-types. Similarly for X-Ray/Radiology: extract ONLY the "Radiology Facility" or "Hospital X-ray" sub-type and ignore "Radiology Office" and "Free Standing" variants. If only a single Lab or X-Ray value exists with no qualifier, extract that single value as-is.
    CRITICAL - DEDUCTIBLE MODIFIER MUST COME FROM THE SAME SUB-ROW: When a cell contains multiple sub-rows (e.g., "Lab Office - $50 Deductible does not apply; Lab Facility - No charge; Radiology Office - $150 Deductible applies; Radiology Facility - $0 Deductible applies"), you MUST extract the cost AND its deductible phrase from the SAME sub-row. Do NOT mix phrases between sub-rows. Example: Lab = "Lab Facility - No charge" → cost=$0, deductible phrase="No charge" (treat as 'Deductible does not apply'). X-Ray = "Radiology Facility - $0/visit Deductible applies" → cost=$0, deductible phrase='Deductible applies'. Do NOT apply the Lab Facility's "No charge / Deductible does not apply" phrase to the X-Ray field.
    Examples: "Lab Facility: No charge" → extract "$0" for lab_services_copay, deductible does NOT apply. "Radiology Facility: $0/visit Deductible applies" → extract "$0" for xray_copay, deductible DOES apply (After Deductible). "Hospital Lab: $15 copay" → extract "$15" for lab_services_copay. "Hospital X-ray: $50 copay" → extract "$50" for xray_copay.

MEDICAL VALUE SELECTION (HOSPITAL/FACILITY PRIORITY):
15b. CRITICAL FOR MULTIPLE FACILITY TYPES: For Inpatient, Outpatient, Lab, X-Ray, and Imaging benefits, if the document lists multiple costs based on the facility type (e.g., "Office" vs "Free-Standing" vs "Hospital" vs "Facility"), you MUST ALWAYS select and extract the "Hospital" or "Facility" value. Completely ignore the costs for Office, Free-Standing, or Ambulatory Surgery Centers. The source SBC may use different terminology like "Hospital Lab", "Hospital X-ray", "Lab Facility", "Radiology Facility", "Hospital", "Facility", etc. Always extract the Hospital/Facility cost for these specific benefits. IMPORTANT: The deductible modifier ('Deductible applies' or 'Deductible does not apply') must also be taken from this same Hospital/Facility sub-row — never from the Office or Free-Standing sub-row.

PHARMACY EXTRACTION RULES:
16. PHARMACY TIER VALUES: Use the "STRUCTURED TABLES" provided in the text for accurate column alignment. You MUST extract ONLY the "In-Network Provider" (or "Participating Provider") value for Tiers 1-5. Do not extract Out-of-Network or Non-Participating values. 

⚠️ CRITICAL FOR PREFERRED VS NON-PREFERRED RETAIL (tier_1 through tier_5 copay and coinsurance ONLY): 
If a pharmacy tier retail cell lists BOTH "Preferred" and "Non-preferred" costs, you MUST ALWAYS extract the LARGER/NON-PREFERRED value and COMPLETELY IGNORE the Preferred value. 

EXAMPLES:
- "Retail Preferred: $10 / Retail Non-preferred: $20" → extract "$20" (the NON-PREFERRED one)
- "Preferred: 20% / Non-preferred: 30%" → extract "30%" (the NON-PREFERRED one)  
- "Generic drugs (Preferred): No Charge, $10 / Generic drugs (Non-preferred): $10, $20" → extract "$20" for this tier (NON-PREFERRED value)
- "Brand drugs (Preferred): $50, $70 / Brand drugs (Non-preferred): $100, $120" → extract "$120" for this tier (NON-PREFERRED value)

⚠️ MANDATORY: When you see TWO values separated by commas in a cell (e.g., "$50, $70" or "No Charge, $10"), the SECOND value is NON-PREFERRED. Extract ONLY the second value.

This rule does NOT apply to specialty_rx_description, tier_4_maximum, or tier_5_maximum (those fields keep both values per Rule 19). 

CRITICAL FOR MULTI-NETWORK CELLS (ALL CARRIERS): If a single cell lists BOTH "Preferred Participating" and "Participating", you MUST extract the pure "Participating" value and IGNORE the "Preferred Participating" value. (e.g., if a cell says "Retail – Preferred Participating – $10 | Participating – $20", extract "$20"). 

CRITICAL FOR ANTHEM PLANS: If there is a "Level 1 Pharmacy" column, DO NOT extract from it; extract from the subsequent "In-Network Provider" column instead. If both retail and home delivery/mail order are present in a cell, extract ONLY the retail value.
    
    CRITICAL FOR PIPE-SEPARATED PHARMACY TABLES (Anthem & similar plans):
    Some plans (e.g., Anthem) use pipe-separated tables with this exact structure:
    | Tier Description | Retail + Home Delivery | In-Network Retail Only | Out-of-Network |
    
    
    For these tables, extract ONLY from the "In-Network Retail Only" column (3rd column in pipes).
    When a cell shows both retail and home delivery like "$90 (retail) and $225 (home delivery)", extract ONLY the retail value "$90".
    When a cell shows "In-Network Retail Only" like "$100/prescription (retail only)", extract "$100".
    
    Example extractions for Anthem Bronze PPO:
    - Tier 1: "$20/prescription (retail only)" → extract "$20"
    - Tier 2: "$100/prescription (retail only)" → extract "$100" (NOT $90 retail or $225 home delivery from same tier)
    - Tier 3: "$170/prescription (retail only)" → extract "$170" (NOT $160 retail or $400 home delivery from same tier)
    - Tier 4: "40% coinsurance up to $500/prescription" → extract "40%" in coinsurance field and "$500" in tier_4_maximum field (set copay to "$0")
    
    CRITICAL: These are EXAMPLES ONLY from specific plans. DO NOT use them as defaults if they do not appear in the file.
    For ANY plan, extract ONLY the ACTUAL values you see in the document's In-Network column. 
    - If Tier 4 shows "40% coinsurance up to $500" → extract "40%" and "$500"
    - If Tier 4 shows "20% coinsurance up to $250" → extract "20%" and "$250"
    - If Tier 4 shows something different → extract that instead
    - If Tier 4 is NOT in the file → do NOT add these example values as defaults
17. PHARMACY TIER COMPLETENESS: You MUST extract ALL pharmacy tiers present in the document (typically Tiers 1-5). Do NOT skip any tier. If a tier row is present in the structured table, extract its value from the In-Network column. Common issue: Some extractors skip Tier 2 or Tier 3. CRITICAL: Count the tiers and ensure all are present in your output. Tiers with values like "$20/Rx", "$100/Rx", "$170/Rx", "40% coinsurance", "$500/Rx" must all be extracted.
17a. CRITICAL FOR PHARMACY COINSURANCE: If a pharmacy tier ONLY shows a percentage (e.g., "20% coinsurance") and no dollar amount, you MUST put the percentage in the COINSURANCE field and "$0" in the COPAY field. NEVER put a standalone percentage like "0%" or "20%" into a pharmacy copay field.
18. MULTIPLE VALUES: If a service lists BOTH a copay AND a coinsurance (e.g., "$35 copayment 50% coinsurance"), you MUST extract both and place them into their respective fields. However, if they are separated by a pipe (e.g., "| $35 copay | 50% coinsurance |"), the second value is Out-of-Network and MUST BE IGNORED. Do not extract the Out-of-Network value.
19. SPECIALTY RX DESCRIPTION: This field captures specialty drug costs from the In-Network column ONLY. Look for specialty drug information in TWO possible formats:
    FORMAT A (Dedicated Specialty Row): Look for a dedicated "Specialty drugs" row in the pharmacy table (usually after Tier 3 or Tier 4). Extract the complete In-Network value including all information.
    FORMAT B (Inline Specialty): When specialty drugs appear inline with each tier (e.g., "Tier 3: Retail: $100 copay | Specialty Drugs: $350 copay"), extract and COMBINE all tier specialty values in order. Example: If Tier 1 has "Specialty: $5", Tier 2 has "Specialty: $150", Tier 3 has "Specialty: $350", Tier 4 has "Specialty: $500", extract as "$5/$150/$350/$500".
    In BOTH formats: Include BOTH percentages AND dollar amounts if both appear. Examples: If In-Network shows "30% (preferred), 50% (non-preferred)", extract as "30%/50%". If it shows "$300/$500", extract as-is. If it shows "Applicable cost as noted above for generic or brand drugs", "Same as above", or "Matches Previous Tiers", extract that exact phrase into specialty_rx_description AND set specialty_mirrors_tiers_1_3 to true. 
    PREFERRED VS NON-PREFERRED SPECIALTY RULES: If the document lists "Specialty drugs" and splits them into "Preferred" and "Non-preferred", you MUST map them as follows ONLY IF the document DOES NOT have explicit Tier 4 and Tier 5 rows (like "Non-preferred brand"):
    - Map "Preferred" Specialty costs to Tier 4.
    - Map "Non-preferred" Specialty costs to Tier 5.
    CRITICAL: If the document ALREADY has explicit Tier 4 (e.g., "$120") and Tier 5 (e.g., "$150") costs, DO NOT overwrite them! Leave Tier 4 and Tier 5 as their explicit values, and just put the specialty costs in specialty_rx_description.
    - If the text says "up to" or "maximum" (e.g., "up to $250" or "$250 maximum copay"), place that value into the corresponding tier_X_maximum field, NOT the regular tier_X_copay field. For example, if it says "up to $250 (preferred) and $500 (non-preferred)", put "$250" in tier_4_maximum and "$500" in tier_5_maximum, and set regular copays to "$0".
    - STILL combine both values (e.g., "20%/40%") into the specialty_rx_description field.
    CRITICAL: If NO specialty drug information exists anywhere in the document, set specialty_rx_description to null (do NOT default to "$0" or use the deductible amount). Extract ONLY from In-Network column, NEVER from Out-of-Network or Limitations columns. Do not overwrite or remove existing copay/coinsurance values from Tiers 1-5; they must coexist with the specialty drug values.

BOOLEAN AND IDENTIFICATION RULES:
20. BOOLEAN FIELDS: For HDHP, Open Access, Offers Tier 1A Benefit, Specialty Mirrors Tiers 1-3, and Out Of Network Coverage, return true or false.
21. BLANK OR MISSING VALUES: If a service is not mentioned or is completely blank for a specific field:
    - For COINSURANCE fields: return "0%".
    - For COPAY, DEDUCTIBLE, and OOP MAX fields: return "$0".
    - Do NOT return null for these numeric cost fields.
22. CARRIER: The carrier is the HEALTH INSURANCE COMPANY that underwrites the plan (e.g., Cigna, Aetna, UnitedHealthcare, BlueCross, Blue Shield, Anthem, Kaiser Permanente, Humana). It is NEVER the employer, company, or group that offers the plan to employees. Look for carrier names in headers, footers, or document titles. Common patterns: "UnitedHealthcare", "Cigna", "Aetna", "Anthem", "Blue Cross", "Kaiser".
23. PLAN_NAME: Extract the actual insurance plan name using the following dynamic strategy:
    STEP 1 - FIND THE HEADER LINE: Look at the very first lines of the document (before any tables or body text). SBC documents always start with a header line in this exact format:
        "[Carrier Name] [Plan Name] Coverage for: [Individual/Family] Plan Type: [PPO/HMO/EPO/POS/HDHP]"
        Example: "Aetna Open Access Managed Choice - NY OA MC 2000/100% Coverage for: Individual + Family Plan Type: POS"
        In this example, the plan name is everything between the carrier name and "Coverage for:" → "Open Access Managed Choice - NY OA MC 2000/100%"
        Prefer the short code portion if both a long name and a short code are present (e.g., prefer "NY OA MC 2000/100%" over the full long name).
    STEP 2 - LOOK FOR PLAN CODES: Plan names are typically short codes or product names like:
        "NY OA MC 2000/100%", "OXF-FRE EPO 5B 1500-80", "Choice Plus HDHP 3300", "OAMC HDHP 3500/80", "Bronze PPO 7000"
        These are concise, contain numbers, slashes, dashes, and plan type abbreviations (PPO, HMO, EPO, POS, HDHP, OA, MC).
    REJECT ALL of the following — they are NOT plan names:
        - Glossary or legal text: any sentence containing "balance billing", "coinsurance", "copayment", "deductible", "Glossary", "underlined terms", "see the Glossary" → REJECT
        - Employer/group names: any name containing "INC.", "LLC", "PEO", "CORP", "CORP.", "GROUP", "SERVICES", "SOLUTIONS", "ENTERPRISES", "FOUNDATION", "RIPPLING", "HOLDINGS", "TRUST", "ASSOCIATION" → REJECT
        - OCR logo artifacts: short garbled strings like "vaetna", "Bcigna", "Antheim", "Unitedh" (carrier name with extra/missing letters) → REJECT
        - Boilerplate phrases: "Summary of Benefits", "Summary of Benefits and Coverage", "Important Questions", "What You Will Pay" → REJECT
        - Long sentences (more than 80 characters that are not a plan code) → REJECT
    FALLBACK: If you cannot confidently identify the plan name from the document, return null. Do NOT guess or fabricate.
24. PLAN_SOURCE: Extract the unique plan identifier, member ID prefix, or policy number if available from the document. If not available, return null (the pipeline will use the filename). Do NOT return "PDF" or generic values.

HDHP (High Deductible Health Plan) DETECTION:
- Look ONLY for explicit mentions of "HDHP" or "HSA" in the plan type, plan name, or headers.
- CRITICAL: DO NOT guess HDHP status based on high dollar amounts like $1,400+ deductibles. It MUST explicitly state "HDHP" or "HSA".
- Set hdhp: true only if these exact words are explicitly detected.
"""

        prompt = f"""
        Extract the Summary of Benefits and Coverage data from the following text.
        The text includes raw text and STRUCTURED TABLES. You should heavily rely on the STRUCTURED TABLES for extracting costs.
        Return the structured JSON exactly matching the requested schema.

        {FORMATTING_RULES}

        Document Text:
        {text}
        """

        print(f"  [LLM] Sending {len(text)} chars to GPT-4o for parsing...")
        
        # Retry logic for OpenAI API calls with exponential backoff
        max_retries = 3
        base_delay = 2
        
        for attempt in range(max_retries):
            try:
                completion = self.client.beta.chat.completions.parse(
                    model="gpt-4o-2024-08-06",
                    messages=[
                        {"role": "system", "content": "You are a precise data extraction expert reading SBC health insurance documents. You MUST follow the formatting rules provided by the user exactly."},
                        {"role": "user", "content": prompt}
                    ],
                    response_format=MasterSBCSchema,
                )
                break  # Success, exit retry loop
                
            except Exception as e:
                print(f"  [LLM] Attempt {attempt + 1} failed: {str(e)}")
                
                if attempt == max_retries - 1:  # Last attempt
                    print(f"  [LLM] All {max_retries} attempts failed, raising exception")
                    raise e
                
                # Wait before retrying with exponential backoff
                delay = base_delay * (2 ** attempt)
                print(f"  [LLM] Retrying in {delay} seconds...")
                import time
                time.sleep(delay)
        
        result = completion.choices[0].message.parsed

        # Normalize plan name and pharmacy tiers using text-backed rules
        result.plan_information.plan_name = self._extract_plan_name_from_text(
            text,
            result.plan_information.plan_name,
            result.plan_information.carrier or ''
        )
        for tier_num in range(1, 6):
            for field_suffix in ('copay', 'coinsurance'):
                tier_field = f'tier_{tier_num}_{field_suffix}'
                current_value = getattr(result.pharmacy, tier_field, None)
                normalized = self._normalize_pharmacy_tier(current_value)
                setattr(result.pharmacy, tier_field, normalized)

        # Post-process office visit coinsurance: prefer explicit copay when present
        # and only keep a coinsurance % for PCP/Specialist if a percent is explicitly
        # present near the corresponding keyword in the source text. This avoids
        # pulling a generic "20%" from nearby imaging/facility rows into the
        # ──────────────────────────────────────────────────────────────────────
        # CRITICAL: Pharmacy and Medical rules moved to rules_engine.py
        
        def fix_medical_modifiers(section, copay_mod_field, coinsurance_mod_field, service_name):
            """Ensure copay_modifier and coinsurance_modifier match if needed (we'll let rules_engine handle the real logic)"""
            pass

        print(f"  [LLM] Response received")
        print(f"    Carrier: {result.plan_information.carrier}")
        print(f"    Plan: {result.plan_information.plan_name}")
        print(f"    Type: {result.plan_information.plan_type}")
        print(f"    PC Copay: {result.office_visits.primary_care_copay}")
        print(f"    ER Copay: {result.hospital_surgical.er_copay}")
        print(f"    Tier 1: {result.pharmacy.tier_1_copay}")
        
        # ──────────────────────────────────────────────────────────────────────
        # POST-PROCESSING: Inpatient Copay Per-Day Rule
        # Business rule: if raw text says "X/day" or "per day" near the inpatient
        # section, multiply the extracted copay by 3 (3-day max exposure standard)
        # ──────────────────────────────────────────────────────────────────────
        inp_copay = result.hospital_surgical.inpatient_copay
        if inp_copay:
            # Collect text around the inpatient / hospital stay context
            inpatient_section = ""
            idx = text.lower().find('inpatient')
            if idx != -1:
                inpatient_section += text[max(0, idx - 100):idx + 600]
            idx2 = text.lower().find('hospital stay')
            if idx2 != -1:
                inpatient_section += text[max(0, idx2 - 100):idx2 + 600]
            
            has_per_day = bool(re.search(
                r'(copay/day|per\s+day|/day|\b\d+\s+days?\b)',
                inpatient_section, re.IGNORECASE
            ))
            
            if has_per_day:
                dollar_match = re.search(r'\$([\d,]+)', str(inp_copay))
                if dollar_match:
                    amount = int(dollar_match.group(1).replace(',', ''))
                    new_amount = amount * 3
                    result.hospital_surgical.inpatient_copay = f"${new_amount:,}"
                    # Mark this copay as "per-day calculated" so validation doesn't delete it
                    result.hospital_surgical.inpatient_copay_is_per_day = True
                    print(f"  [FIX] Inpatient Copay: Per-day detected -> ${amount} x 3 = ${new_amount:,}")
            else:
                print(f"  [OK] Inpatient Copay: No per-day charge, keeping '{inp_copay}'")

        # ──────────────────────────────────────────────────────────────────────
        # CRITICAL: Pharmacy and Medical rules moved to rules_engine.py
        
        # ──────────────────────────────────────────────────────────────────────
        # ─── POST-PROCESSING: COINSURANCE-ONLY PLANS VALIDATION ─────────────
        # ──────────────────────────────────────────────────────────────────────
        # Fix cases where document shows coinsurance but AI extracted as copays
        
        print(f"\n  [VALIDATION] Checking for coinsurance-only plan patterns...")
        
        def has_consistent_coinsurance_pattern(text):
            """Check if document consistently shows coinsurance (e.g., 10%, 20%) without copay amounts"""
            # Look for patterns like "10% coinsurance" appearing multiple times
            coinsurance_matches = re.findall(r'\b(\d+)%\s*coinsurance', text, re.IGNORECASE)
            dollar_matches = re.findall(r'\$\d+\s*(?:copay|copayment)', text, re.IGNORECASE)
            
            # If we see repeated coinsurance percentages (3+ times) but copays are much less frequent
            return len(coinsurance_matches) >= 3 and len(coinsurance_matches) > len(dollar_matches) * 2
        
        if has_consistent_coinsurance_pattern(text):
            print(f"    [VALIDATION] Detected coinsurance-only pattern - validating extractions...")
            
            # Common coinsurance percentages found in text
            coinsurance_values = re.findall(r'\b(\d+)%', text)
            most_common_coinsurance = None
            if coinsurance_values:
                from collections import Counter
                counter = Counter(coinsurance_values)
                most_common_coinsurance = f"{counter.most_common(1)[0][0]}%"
                print(f"    [VALIDATION] Most common coinsurance: {most_common_coinsurance}")
            
            # Validate office visits
            office = result.office_visits
            if office.primary_care_copay == "$0" and office.primary_care_coinsurance in [None, "0%"]:
                if most_common_coinsurance:
                    office.primary_care_coinsurance = most_common_coinsurance
                    print(f"    [FIX] Primary Care: Set coinsurance to {most_common_coinsurance}")
            
            if office.specialist_copay == "$0" and office.specialist_coinsurance in [None, "0%"]:
                if most_common_coinsurance:
                    office.specialist_coinsurance = most_common_coinsurance
                    print(f"    [FIX] Specialist: Set coinsurance to {most_common_coinsurance}")
            
            # Validate hospital services
            hospital = result.hospital_surgical
            for service_name, copay_field, coinsurance_field in [
                ("Inpatient", "inpatient_copay", "inpatient_coinsurance"),
                ("Outpatient", "op_hospital_copay", "op_hospital_coinsurance"),
                ("Emergency Room", "er_copay", "er_coinsurance")
            ]:
                copay_val = getattr(hospital, copay_field, None)
                coinsurance_val = getattr(hospital, coinsurance_field, None)
                
                if copay_val == "$0" and coinsurance_val in [None, "0%"]:
                    if most_common_coinsurance:
                        setattr(hospital, coinsurance_field, most_common_coinsurance)
                        print(f"    [FIX] {service_name}: Set coinsurance to {most_common_coinsurance}")
            
            # Validate urgent care & imaging
            imaging = result.urgent_care_labs_imaging
            for service_name, copay_field, coinsurance_field in [
                ("Urgent Care", "urgent_care_copay", "urgent_care_coinsurance"),
                ("Lab Services", "lab_services_copay", "lab_services_coinsurance"),
                ("X-Ray", "xray_copay", "xray_coinsurance"),
                ("Medical Imaging", "medical_imaging_copay", "medical_imaging_coinsurance")
            ]:
                copay_val = getattr(imaging, copay_field, None)
                coinsurance_val = getattr(imaging, coinsurance_field, None)
                
                if copay_val == "$0" and coinsurance_val in [None, "0%"]:
                    if most_common_coinsurance:
                        setattr(imaging, coinsurance_field, most_common_coinsurance)
                        print(f"    [FIX] {service_name}: Set coinsurance to {most_common_coinsurance}")
        
        # ──────────────────────────────────────────────────────────────────────
        # ─── POST-PROCESSING: DEDUCTIBLE TYPE AND STATUS FIELD POPULATION ───
        # ──────────────────────────────────────────────────────────────────────
        # Fix missing deductible_type and populate deductible_status fields
        
        print(f"\n  [POST] Setting deductible type and status fields...")
        
        # Determine deductible type based on deductible amounts
        deductibles = result.deductibles_and_coinsurance
        individual_ded = deductibles.individual_deductible or "$0"
        family_ded = deductibles.family_deductible or "$0"
        
        if "the overall family deductible must be met before the plan begins to pay" in text.lower():
            deductibles.deductible_type = "True Individual Family"
        else:
            deductibles.deductible_type = "Embedded - Traditional Style"
        
        print(f"    [DEDUCTIBLE] Type set to: {deductibles.deductible_type}")
        
        # Determine default deductible status based on plan structure
        is_no_deductible = (individual_ded == "$0")
        hdhp = result.plan_information.hdhp
        
        if is_no_deductible:
            default_status = "No"  # Deductible doesn't apply
        elif hdhp:
            default_status = "Yes"  # HDHP - deductible applies to most services
        else:
            default_status = "No"   # Traditional plan - many services waived
        
        print(f"    [STATUS] Default deductible status: {default_status}")
        
        # Populate deductible_status fields for all sections
        def set_deductible_status(section, field_pairs):
            for copay_field, coinsurance_field in field_pairs:
                # Set copay deductible status
                copay_status_field = copay_field + '_deductible_status'
                if hasattr(section, copay_status_field) and getattr(section, copay_status_field) is None:
                    modifier = getattr(section, copay_field + '_modifier', None)
                    if modifier and "Deductible Waived" in str(modifier):
                        setattr(section, copay_status_field, "No")
                    elif modifier and "After Deductible" in str(modifier):
                        setattr(section, copay_status_field, "Yes")
                    else:
                        setattr(section, copay_status_field, default_status)
                
                # Set coinsurance deductible status
                coinsurance_status_field = coinsurance_field + '_deductible_status'
                if hasattr(section, coinsurance_status_field) and getattr(section, coinsurance_status_field) is None:
                    modifier = getattr(section, coinsurance_field + '_modifier', None)
                    if modifier and "Deductible Waived" in str(modifier):
                        setattr(section, coinsurance_status_field, "No")
                    elif modifier and "After Deductible" in str(modifier):
                        setattr(section, coinsurance_status_field, "Yes")
                    else:
                        setattr(section, coinsurance_status_field, default_status)
        
        # Apply to office visits
        set_deductible_status(result.office_visits, [
            ('primary_care_copay', 'primary_care_coinsurance'),
            ('specialist_copay', 'specialist_coinsurance')
        ])
        
        # Apply to hospital surgical  
        set_deductible_status(result.hospital_surgical, [
            ('inpatient_copay', 'inpatient_coinsurance'),
            ('op_hospital_copay', 'op_hospital_coinsurance'),
            ('er_copay', 'er_coinsurance')
        ])
        
        # Apply to urgent care, labs & imaging
        set_deductible_status(result.urgent_care_labs_imaging, [
            ('urgent_care_copay', 'urgent_care_coinsurance'),
            ('lab_services_copay', 'lab_services_coinsurance'),
            ('xray_copay', 'xray_coinsurance'),
            ('medical_imaging_copay', 'medical_imaging_coinsurance')
        ])
        
        # Special handling for pharmacy deductible status
        pharmacy = result.pharmacy
        if hasattr(pharmacy, 'pharmacy_deductible_deductible_status') and pharmacy.pharmacy_deductible_deductible_status is None:
            pharmacy_ded = pharmacy.pharmacy_deductible or "$0"
            pharmacy.pharmacy_deductible_deductible_status = "Yes" if pharmacy_ded != "$0" else "No"
        
        # Pharmacy tiers deductible status
        for tier_num in range(1, 6):
            for field_type in ['copay', 'coinsurance']:
                status_field = f'tier_{tier_num}_{field_type}_deductible_status'
                if hasattr(pharmacy, status_field) and getattr(pharmacy, status_field) is None:
                    modifier_field = f'tier_{tier_num}_{field_type}_modifier'
                    modifier = getattr(pharmacy, modifier_field, None)
                    if modifier and "Deductible Waived" in str(modifier):
                        setattr(pharmacy, status_field, "No")
                    elif modifier and "After" in str(modifier):
                        setattr(pharmacy, status_field, "Yes")
                    else:
                        setattr(pharmacy, status_field, default_status)
        
        print(f"  [POST] Deductible type and status fields populated")
        
        print(f"  [POST] All post-processing validation complete")
        
        return result
